"""Quiz export functionality for different formats (PDF, DOCX, JSON)."""

import json
from io import BytesIO
from datetime import datetime
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
from reportlab.lib import colors
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


class QuizExporter:
    """Export quizzes to various formats (PDF, DOCX, JSON)."""
    
    def __init__(self):
        """Initialize the exporter."""
        pass
    
    def export_to_json(self, quiz_text: str, text: str, annotations: dict, tag_type: str) -> str:
        """
        Export quiz to JSON format.
        
        Args:
            quiz_text: The formatted quiz text
            text: Original text
            annotations: Annotations dictionary
            tag_type: Tag type (5W, Thesis, etc.)
            
        Returns:
            JSON string
        """
        quiz_data = {
            "text": text,
            "annotations": annotations,
            "quiz": quiz_text,
            "tag_type": tag_type,
            "generated_at": datetime.now().isoformat()
        }
        return json.dumps(quiz_data, indent=2, ensure_ascii=False)
    
    def export_to_pdf(self, structured_quiz: list, tag_type: str, original_text: str = "") -> bytes:
        """
        Export quiz to PDF format.
        
        Args:
            structured_quiz: List of structured questions
            tag_type: Tag type (5W, Thesis, etc.)
            original_text: Original text (optional, for context)
            
        Returns:
            PDF file as bytes
        """
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=0.75*inch, bottomMargin=0.75*inch)
        story = []
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            textColor=colors.HexColor('#1f77b4'),
            spaceAfter=12,
            alignment=1  # Center
        )
        
        question_style = ParagraphStyle(
            'QuestionStyle',
            parent=styles['Normal'],
            fontSize=12,
            textColor=colors.black,
            spaceAfter=10,
            fontName='Helvetica-Bold'
        )
        
        option_style = ParagraphStyle(
            'OptionStyle',
            parent=styles['Normal'],
            fontSize=11,
            leftIndent=20,
            spaceAfter=6
        )
        
        answer_style = ParagraphStyle(
            'AnswerStyle',
            parent=styles['Normal'],
            fontSize=10,
            textColor=colors.HexColor('#666666'),
            leftIndent=20,
            spaceAfter=6
        )
        
        # Header
        story.append(Paragraph("Linda - AI Assessment Platform", title_style))
        story.append(Paragraph(f"Quiz Type: {tag_type}", styles['Normal']))
        story.append(Paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}", styles['Normal']))
        story.append(Spacer(1, 0.3*inch))
        
        # Questions
        for q in structured_quiz:
            # Question number and text
            story.append(Paragraph(f"<b>Question {q['number']}:</b> {q['text']}", question_style))
            story.append(Spacer(1, 0.1*inch))
            
            if q['type'] == 'multiple_choice':
                # Options
                for opt in q['options']:
                    is_correct = opt['letter'] == q['correct_answer']
                    marker = "✓" if is_correct else "○"
                    story.append(Paragraph(f"{marker} {opt['letter']}) {opt['text']}", option_style))
                story.append(Spacer(1, 0.15*inch))
            else:
                # Open-ended
                story.append(Paragraph("<i>Open-ended question - student provides written answer</i>", option_style))
                story.append(Spacer(1, 0.1*inch))
                story.append(Paragraph(f"<b>Sample Answer:</b> {q['correct_answer']}", answer_style))
                story.append(Spacer(1, 0.15*inch))
        
        # Answer key on separate page
        story.append(PageBreak())
        story.append(Paragraph("Answer Key", title_style))
        story.append(Spacer(1, 0.2*inch))
        
        answer_data = [["Question", "Correct Answer"]]
        for q in structured_quiz:
            if q['type'] == 'multiple_choice':
                answer_text = f"{q['correct_answer']}"
            else:
                answer_text = "Open-ended (see sample)"
            answer_data.append([f"Q{q['number']}", answer_text])
        
        answer_table = Table(answer_data, colWidths=[1.5*inch, 4*inch])
        answer_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1f77b4')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        story.append(answer_table)
        
        # Build PDF
        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()
    
    def export_to_docx(self, structured_quiz: list, tag_type: str, original_text: str = "") -> bytes:
        """
        Export quiz to DOCX (Word) format.
        
        Args:
            structured_quiz: List of structured questions
            tag_type: Tag type (5W, Thesis, etc.)
            original_text: Original text (optional, for context)
            
        Returns:
            DOCX file as bytes
        """
        doc = Document()
        
        # Header
        title = doc.add_heading('Linda - AI Assessment Platform', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Metadata
        meta = doc.add_paragraph()
        meta.add_run(f'Quiz Type: {tag_type}\n').bold = True
        meta.add_run(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M")}\n')
        meta.add_run('_' * 80)
        
        doc.add_paragraph()  # Spacing
        
        # Questions
        for q in structured_quiz:
            # Question heading
            q_heading = doc.add_heading(f"Question {q['number']}", level=2)
            
            # Question text
            q_para = doc.add_paragraph()
            q_run = q_para.add_run(q['text'])
            q_run.font.size = Pt(12)
            
            doc.add_paragraph()  # Spacing
            
            if q['type'] == 'multiple_choice':
                # Options
                for opt in q['options']:
                    is_correct = opt['letter'] == q['correct_answer']
                    opt_para = doc.add_paragraph(style='List Number')
                    opt_run = opt_para.add_run(f"{opt['letter']}) {opt['text']}")
                    
                    if is_correct:
                        opt_run.font.color.rgb = RGBColor(0, 128, 0)  # Green for correct
                        opt_run.bold = True
                        # Add checkmark
                        opt_para.add_run(' ✓').font.color.rgb = RGBColor(0, 128, 0)
            else:
                # Open-ended
                doc.add_paragraph('Open-ended question - student provides written answer', style='List Bullet')
                
                answer_para = doc.add_paragraph()
                answer_para.add_run('Sample Answer: ').bold = True
                answer_para.add_run(q['correct_answer']).italic = True
            
            doc.add_paragraph()  # Spacing between questions
        
        # Answer key on new page
        doc.add_page_break()
        answer_heading = doc.add_heading('Answer Key', 0)
        answer_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Answer table
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Light Grid Accent 1'
        
        # Header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Question'
        hdr_cells[1].text = 'Correct Answer'
        
        # Make header bold
        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        # Data rows
        for q in structured_quiz:
            row_cells = table.add_row().cells
            row_cells[0].text = f"Q{q['number']}"
            
            if q['type'] == 'multiple_choice':
                row_cells[1].text = f"{q['correct_answer']}"
            else:
                row_cells[1].text = "Open-ended (see sample answer above)"
        
        # Save to buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

